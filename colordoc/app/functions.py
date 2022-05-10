from docx import Document
from docx.shared import RGBColor

def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment

def colorfile(f):
    docOut = Document()
    docIn = Document('media/'+f) 
    # print(f.name)
    for p in docIn.paragraphs: 
        paragraph = docOut.add_paragraph()
        # get_para_data(docOut,paragraph)
        for run in p.runs: 
            # runO.style = run.style
        

            # print(run.text, run.italic, run.bold)
            if run.italic: 
                runO = paragraph.add_run(run.text)
                runO.font.color.rgb = RGBColor(255, 0, 0)
                runO.italic = run.italic
                
                

            elif run.bold:
                runO = paragraph.add_run(run.text)
                runO.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
                runO.bold = run.bold


            elif run.underline:
                runO = paragraph.add_run(run.text)
                runO.font.color.rgb = RGBColor.from_string('0000FF')
                runO.underline = run.underline

            else:
                runO = paragraph.add_run(run.text)
                runO.style = run.style
                runO.font.color.rgb = run.font.color.rgb
                paragraph.paragraph_format.alignment = p.paragraph_format.alignment
                



        docOut.save('media/out/'+f)
    return 'media/out/'+f
